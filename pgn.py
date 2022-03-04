# pylint: disable=import-error
# pylint: disable=protected-access
# pylint: disable=too-many-locals
# pylint: disable=too-many-statements
"""functions for pgn mgmt, and docx generation"""

import io
import os
import os.path
import re
#import sys


import warnings

import chess
import chess.pgn
import numpy as np
import pandas as pd

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import Inches, Mm, Pt, RGBColor
from docx.oxml.ns import qn

import chessboard as cb
import eco

warnings.simplefilter(action='ignore', category=FutureWarning)


def get_pgnfile_names_from_dir(pgn_dir='PGN/', ext='.pgn') -> list:
    """Return a python list with filenames
       from a given directory and given extension '.pgn' '.PGN'"""
    file_names_list = []
    for file in os.listdir(pgn_dir):
        if file.endswith(ext) or file.endswith(ext.upper()):
            if os.path.isfile(os.path.join(pgn_dir, file)):
                file_names_list.append(os.path.join(pgn_dir, file))
    return sorted(file_names_list)


def get_games_from_pgnfile(file_name: str) -> pd.DataFrame:
    """Return a DataFrame with all games of the file_name, incl. headers and pgn game notation."""
    with open(file_name, "r", encoding='utf-8') as pgn_file:
        games_df = pd.DataFrame()
        # iterate over all games of a file
        while True:
            try:
                game = chess.pgn.read_game(pgn_file)
                if game is None:
                    break
            except BaseException as err:
                print(f"Unexpected {err=}, {type(err)=}")
                print('at file:', file_name, 'with')
                print(game)
                continue

            try:
                game_dict = dict(game.headers)
                game_dict["pgn"] = game.board().variation_san(
                    game.mainline_moves())
                game_dict['file'] = file_name
                games_df = games_df.append(game_dict, ignore_index=True)
            except BaseException as err:
                print(f"Unexpected {err=}, {type(err)=}")
                print('at file:', file_name, 'with')
                print(game)
                continue
    return games_df


def prep_ttfboards_from_pgn(pgn_str: str) -> pd.DataFrame:
    """Return at each row full move info: W & B SAN string and W & B TTF board string"""
    # start chess board
    board = chess.Board()
    game = chess.pgn.read_game(io.StringIO(pgn_str))

    # prep for half moves
    half_moves_df = pd.DataFrame()
    for move in game.mainline_moves():
        # dict of half move infos
        move_dict = {}

        move = chess.Move.from_uci(move.uci())
        move_dict['FMVN'] = int(board.fullmove_number)
        move_dict['SAN'] = board.san(move)

        # fromto for later markup
        move_dict['sq_from'] = move.uci()[:2]
        move_dict['sq_to'] = move.uci()[2:4]
        # site to move chess.WHITE or chess.BLACK
        move_dict['player'] = board.turn
        if chess.WHITE == board.turn:
            move_dict['mv_san_str'] = \
                str(move_dict['FMVN']) + '. ' + \
                move_dict['SAN'] + ' ... '
        else:
            move_dict['mv_san_str'] = \
                str(move_dict['FMVN']) + '. ' + \
                ' ... ' + move_dict['SAN']

        board.push(move)
        sq_check = ''
        if board.is_check():
            sq_check = chess.square_name(board.king(board.turn))
        move_dict['sq_check'] = sq_check

        move_dict['board_arr'] = cb.board2arr(board)
        # print(move_dict)
        half_moves_df = half_moves_df.append(move_dict, ignore_index=True)

    half_moves_df = half_moves_df.astype({'FMVN': np.uint8})
    # print(half_moves_df)

    # half moves to full moves data
    # for direct use at document creation
    full_moves_df = pd.DataFrame()
    full_move_dict = {}
    for _, hmv in half_moves_df.iterrows():
        full_move_dict['FMVN'] = int(hmv['FMVN'])
        if chess.WHITE == hmv['player']:
            full_move_dict['w_hmv_str'] = hmv['mv_san_str']
            # depending on marked squares
            # change the boards' ttf representation
            full_move_dict['w_sq_from'] = hmv['sq_from']
            full_move_dict['w_sq_to'] = hmv['sq_to']
            full_move_dict['w_sq_check'] = hmv['sq_check']
            full_move_dict['w_board_ttf'] = cb.arr2ttf(hmv['board_arr'])
        else:
            #full_move_dict['b_hmv_str'] = full_move_dict['w_hmv_str'][:-5] + ' ' +  hmv['SAN']
            full_move_dict['b_hmv_str'] = hmv['mv_san_str']
            # depending on marked squares
            # change the boards' ttf representation
            full_move_dict['b_sq_from'] = hmv['sq_from']
            full_move_dict['b_sq_to'] = hmv['sq_to']
            full_move_dict['b_sq_check'] = hmv['sq_check']
            full_move_dict['b_board_ttf'] = cb.arr2ttf(hmv['board_arr'])
            full_moves_df = full_moves_df.append(
                full_move_dict, ignore_index=True)
    full_moves_df = full_moves_df.astype({'FMVN': np.uint8})
    return full_moves_df


def gen_document_from_game(game_dict: dict,
                           eco_dict: dict,
                           ttf_font_name='Chess Merida') -> Document:
    """Return a docx.Document Din A4 with the chess diagrams for a given game_dict"""

    # if ttf_font_name not in cb.TTF_dict.keys():
    #     print(f'You choose TTF {ttf_font_name},
    #     which is not is my list of fonts to chose from:')
    #     for font in cb.TTF_dict.keys():
    #         print(f'  - {font] with {cb.TTF_dict[font]}')
    #     print('Please install the font you want.')
    #     print('I will continue docx generation with TTF Chess Merida')

    doc = Document()

    #  set to A4 --------------------------------------
    # see
    #   https://stackoverflow.com/questions/43724030
    #   /how-to-change-page-size-to-a4-in-python-docx
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(30)
    section.right_margin = Mm(25)

    section.header_height = Inches(0.2)
    section.bottom_height = Inches(0.2)

    #section.top_margin = Inches(0.4)
    #section.bottom_margin = Inches(0.4)

    #section.header_distance = Inches(0.1)
    #section.footer_distance = Inches(0.1)
    #  set to A4 --------------------------------------

    # doc header --------------------------------------
    header = doc.sections[0].header
    #header.bottom_margin = Inches(0.2)
    head = header.paragraphs[0]
    head.text = f"{game_dict['Date'].replace('.','-')} " + \
        f"{game_dict['Event']}, {game_dict['Site']}\n" + \
        f"{game_dict['White']} vs. {game_dict['Black']}   " + \
        f"{game_dict['Result']}"
    # doc header --------------------------------------

    # doc footer --------------------------------------
    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        return element.set(ns.qn(name), value)

    def add_page_number(paragraph):
        page_run = paragraph.add_run()
        elem_t1 = create_element('w:t')
        create_attribute(elem_t1, 'xml:space', 'preserve')
        elem_t1.text = 'Page '
        page_run._r.append(elem_t1)

        page_num_run = paragraph.add_run()

        fld_char1 = create_element('w:fldChar')
        create_attribute(fld_char1, 'w:fldCharType', 'begin')

        instr_text = create_element('w:instrText')
        create_attribute(instr_text, 'xml:space', 'preserve')
        instr_text.text = "PAGE"

        fld_char2 = create_element('w:fldChar')
        create_attribute(fld_char2, 'w:fldCharType', 'end')

        page_num_run._r.append(fld_char1)
        page_num_run._r.append(instr_text)
        page_num_run._r.append(fld_char2)

        of_run = paragraph.add_run()
        elem_t2 = create_element('w:t')
        create_attribute(elem_t2, 'xml:space', 'preserve')
        elem_t2.text = ' of '
        of_run._r.append(elem_t2)

        fld_char3 = create_element('w:fldChar')
        create_attribute(fld_char3, 'w:fldCharType', 'begin')

        instr_text2 = create_element('w:instrText')
        create_attribute(instr_text2, 'xml:space', 'preserve')
        instr_text2.text = "NUMPAGES"

        fld_char4 = create_element('w:fldChar')
        create_attribute(fld_char4, 'w:fldCharType', 'end')

        num_pages_run = paragraph.add_run()
        num_pages_run._r.append(fld_char3)
        num_pages_run._r.append(instr_text2)
        num_pages_run._r.append(fld_char4)

    add_page_number(doc.sections[0].footer.paragraphs[0])
    doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    #  doc footer --------------------------------------

    # first page of the booklet
    out_str = ''
    for key in game_dict.keys():
        # if (key != 'file' and key != 'pgn'):
        if key not in ('file', 'pgn'):
            out_str = out_str + \
                '[{k}] \"{v}\"\n'.format(k=key, v=game_dict[key])

    out_str = out_str + '\n'
    out_str = out_str + f"{game_dict['pgn']}  {game_dict['Result']}\n"
    doc.add_paragraph(out_str)

    # some words about the game's ECO
    if not eco_dict:
        # no eco section to print, e.g no eco found
        pass
    else:
        if 'eco' not in eco_dict.keys():
            # no eco section to print, e.g no eco found
            pass
        else:
            eco_txt = \
                f"{eco_dict['eco']} - {eco_dict['title']}\n" + \
                f"{eco_dict['pgn']} \n"

            doc.add_paragraph(eco_txt)

            eco_board = chess.Board(eco_dict['fen'])
            eco_tbl = doc.add_table(2, 1)
            eco_row = eco_tbl.rows[0]
            eco_row.cells[0].text = cb.board2ttf(eco_board)[:-1]
            eco_cell_paragraph = eco_row.cells[0].paragraphs[0]
            eco_cell_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            eco_cell_paragraph.paragraph_format.keep_with_next = True
            eco_cell_paragraph.paragraph_format.space_before = Pt(0)
            eco_cell_paragraph.paragraph_format.space_after = Pt(0)
            eco_run = eco_cell_paragraph.runs
            eco_fnt = eco_run[0].font
            eco_fnt.name = ttf_font_name
            eco_fnt.size = Pt(20)

            eco_row = eco_tbl.rows[1]
            eco_row.cells[0].text = eco_dict['last_ply']
            eco_row.cells[0].paragraphs[0].style.font.name = 'Verdana'
            eco_row.cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)

    doc.add_page_break()

    #  PGN diagramms --------------------------------------
    # the PGN data for diagram genration
    boards_df = prep_ttfboards_from_pgn(game_dict['pgn'])

    def gen_brd_cell(cell,
                     ttf_str: str,
                     sq_check: str,
                     sq_from: str,
                     sq_to: str):
        # make parts according the squares to mark
        ttf_parts_df = cb.divide_ttf_str(ttf_str,
                                         sq_check,
                                         sq_from,
                                         sq_to)

        for index, ttf_part in enumerate(ttf_parts_df['part']):
            brd_cell_paragraph = cell.paragraphs[0]
            brd_cell_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            brd_cell_paragraph.paragraph_format.keep_with_next = True
            brd_cell_paragraph.paragraph_format.space_before = Pt(0)
            brd_cell_paragraph.paragraph_format.space_after = Pt(0)
            run = brd_cell_paragraph.add_run(ttf_part)
            run.font.name = ttf_font_name

            # black
            if ttf_parts_df.iloc[index]['type'] == 'norm':
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                run.font.size = Pt(16)

            # red for king in check
            if ttf_parts_df.iloc[index]['type'] == 'sq_check':
                #run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
                #run.font.size = Pt(16)
                tag = run._r
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'fc3535')
                run.font.size = Pt(16)
                tag.rPr.append(shd)

            # lightgreen for sq_from and sq_to squares
            if ttf_parts_df.iloc[index]['type'] in ('sq_from', 'sq_to'):
                tag = run._r
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'cddba7')
                run.font.size = Pt(16)
                tag.rPr.append(shd)

    boards_tbl = doc.add_table(2*len(boards_df), 2)
    for index, fmv in boards_df.iterrows():

        brd_row = boards_tbl.rows[2*index]

        # the board diagrams
        gen_brd_cell(brd_row.cells[0],
                     fmv['w_board_ttf'][:-1],
                     fmv['w_sq_check'],
                     fmv['w_sq_from'],
                     fmv['w_sq_to'])

        gen_brd_cell(brd_row.cells[1],
                     fmv['b_board_ttf'][:-1],
                     fmv['b_sq_check'],
                     fmv['b_sq_from'],
                     fmv['b_sq_to'])

        # the SAN below the board diagrams
        # if last move, add result
        if index == len(boards_df)-1:
            if len(fmv['b_hmv_str']) == 0:
                fmv['w_hmv_str'] = fmv['w_hmv_str'] + \
                    '   ' + game_dict['Result']
            else:
                fmv['b_hmv_str'] = fmv['b_hmv_str'] + \
                    '   ' + game_dict['Result']

        brd_row = boards_tbl.rows[2*index+1]

        brd_row.cells[0].text = fmv['w_hmv_str']
        brd_row.cells[0].paragraphs[0].style.font.name = 'Verdana'
        brd_row.cells[0].paragraphs[0].style.font.size = Pt(9)
        brd_row.cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)

        brd_row.cells[1].text = fmv['b_hmv_str']
        brd_row.cells[1].paragraphs[0].style.font.name = 'Verdana'
        brd_row.cells[0].paragraphs[0].style.font.size = Pt(9)
        brd_row.cells[1].paragraphs[0].paragraph_format.space_after = Pt(0)

    #  PGN diagramms --------------------------------------
    return doc


def get_incremented_filename(filename: str) -> str:
    """Return the given filename if exists with an increment"""
    name, ext = os.path.splitext(filename)
    seq = 0
    # continue from existing sequence number if any
    rex = re.search(r"^(.*)-(\d+)$", name)
    if rex:
        name = rex[1]
        seq = int(rex[2])

    while os.path.exists(filename):
        seq += 1
        filename = f"{name}-{seq}{ext}"
    return filename


def store_document(doc: Document, file_name: str) -> dict:
    """Return a dict{'done' : True,
    'file_name' : <file_name>} after
    Document is stored at 'file_name'"""
    file_name = get_incremented_filename(file_name)
    doc.save(file_name)
    # check that file name ist stored
    return({'done': os.path.exists(file_name),
            'file_name': file_name})


def main():
    """some test for the pgn.py"""
    ##################################################
    # this is for testing only
    # you need to create at your working directory
    # a directory PGN/, structured as the repo's PGN dir
    # or
    # change the code at
    #   pgn.main pgn_dir to whatever you need
    # the line here :
    ##################################################
    pgn_name = 'test/pgn/test_do_not_change.pgn'

    # start to get the games out of one pgn file
    games_df = get_games_from_pgnfile(pgn_name)

    one_game_dict = games_df.iloc[1].to_dict()

    eco_result_dict = {}
    if 'ECO' in one_game_dict.keys():
        eco_result_dict = \
            eco.new_get_eco_data_for(eco=one_game_dict['ECO'],
                                     pgn=one_game_dict['pgn'])
    else:
        eco_result_dict = \
            eco.new_get_eco_data_for(eco='',
                                     pgn=one_game_dict['pgn'])

    my_doc = gen_document_from_game(one_game_dict,
                                    eco_result_dict)

    # fn fix for lichess pgn files
    event = one_game_dict['Event'].replace(
        '/', '_').replace(':', '_').replace('.', '-')
    site = one_game_dict['Site'].replace(
        '/', '_').replace(':', '_').replace('.', '-')

    fname = 'test/docx/' + one_game_dict['Date'].replace('.', '-') + \
            '_' + \
            event + '_' + \
            site + '_( ' + \
            one_game_dict['White'] + ' - ' + \
            one_game_dict['Black'] + ' ).docx'
    fname = fname.replace('??', '_')

    # store document
    ret_dict = store_document(my_doc, fname)
    print('stored:', ret_dict['file_name'])

    return()


if __name__ == '__main__':
    main()
